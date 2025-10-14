from langchain.schema.document import Document
from typing import List
import logging
import os


import fitz
import easyocr
import numpy as np
from PIL import Image
import io
import docx


logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

class DocumentLoader:
    """
    Advanced loader supporting PDF & DOCX.
    For PDFs, it extracts text and uses OCR to read text from diagrams/images.
    """

  
  
    def __init__(self, clean: bool = True, use_ocr: bool = True):
        self.clean = clean
        self.use_ocr = use_ocr
        if self.use_ocr:
            logging.info("Initializing OCR model (EasyOCR)... This may take a few moments.")
            self.ocr_reader = easyocr.Reader(['en'])
            logging.info("OCR model initialized.")

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean text by removing extra whitespace and normalizing line breaks."""
        return ' '.join(text.split())


    def _perform_ocr_on_image_bytes(self, image_bytes: bytes, context: str) -> str:
        """Helper function to run OCR on image bytes and return formatted text."""
        try:

            image = Image.open(io.BytesIO(image_bytes))
            img_np = np.array(image)
            
            ocr_results = self.ocr_reader.readtext(img_np)
            
            text_from_image = " ".join([text for _, text, _ in ocr_results])
            if text_from_image.strip():
                return f"\n{context}: {text_from_image}"
            return ""
            
        except Exception as e:
            logging.warning(f"Could not process image for OCR ({context}): {e}")
            return ""


    def load_pdf(self, path: str) -> List[Document]:
        """
        Loads a PDF, extracting text and performing OCR on images page-by-page.
        Each page becomes a separate Document.
        """

        if not os.path.exists(path):
            raise FileNotFoundError(f"PDF file not found: {path}")
        
        logging.info(f"Loading PDF with OCR: {path}")
        pdf_doc = fitz.open(path)
        all_docs = []
        
        for page_num, page in enumerate(pdf_doc, start=1):
            page_text = page.get_text()
            ocr_text = ""
            
            if self.use_ocr:
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list, start=1):
                    xref = img[0]
                    base_image = pdf_doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    context = f"[Diagram Text on Page {page_num}, Image {img_index}]"
                    ocr_text += self._perform_ocr_on_image_bytes(image_bytes, context)


            combined_content = page_text + ocr_text
            if self.clean:
                combined_content = self._clean_text(combined_content)
            
            metadata = {"source": path, "page": page_num}
            page_document = Document(page_content=combined_content, metadata=metadata)
            all_docs.append(page_document)

        logging.info(f"Loaded {len(all_docs)} pages from PDF with OCR data.")
        return all_docs




    def load_docx(self, path: str) -> List[Document]:
        """
        Loads a DOCX, extracting all text and performing OCR on all images.
        The entire file becomes a single Document.
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        logging.info(f"Loading DOCX with OCR: {path}")
        docx_doc = docx.Document(path)
        
        full_text = [para.text for para in docx_doc.paragraphs]
        doc_text = "\n".join(full_text)
        
        ocr_text = ""
        if self.use_ocr:

            image_index = 1
            for rel in docx_doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_bytes = rel.target_part.blob
                    context = f"[Diagram Text, Image {image_index}]"
                    ocr_text += self._perform_ocr_on_image_bytes(image_bytes, context)
                    image_index += 1
        

        combined_content = doc_text + ocr_text
        if self.clean:
            combined_content = self._clean_text(combined_content)
        
        metadata = {"source": path}
        doc = Document(page_content=combined_content, metadata=metadata)
        
        logging.info(f"Loaded DOCX file as a single document with OCR data.")
        return [doc]




    def load_document(self, path: str) -> List[Document]:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            return self.load_pdf(path)
        elif ext == ".docx":
            return self.load_docx(path)
        else:
            raise ValueError(f"Unsupported file format '{ext}'. Only PDF and DOCX are supported.")